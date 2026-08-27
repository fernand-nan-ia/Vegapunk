from pathlib import Path

import pytest

from vegapunk import vault
from vegapunk.extract import ExtractionError, extract_document
from vegapunk.normalize import normalize
from vegapunk import voices


def _docx(path: Path):
    import docx
    d = docx.Document()
    d.add_heading("Laudo de Avaliação", level=1)
    d.add_paragraph("Parágrafo com conteúdo suficiente para passar do mínimo de extração. " * 4)
    d.add_heading("Metodologia", level=2)
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text, t.cell(0, 1).text, t.cell(1, 0).text, t.cell(1, 1).text = "Item", "Valor", "Terreno", "1000"
    d.save(str(path))


def _xlsx(path: Path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active; ws.title = "Custos"
    ws.append(["Item", "Valor"]); ws.append(["Servidor", 120.5]); ws.append(["Domínio", 40])
    ws2 = wb.create_sheet("Vazia")
    for _ in range(20):
        ws.append(["linha de preenchimento para passar dos 200 chars", 1])
    wb.save(str(path))


def test_docx_headings_and_tables(tmp_path):
    p = tmp_path / "laudo_final.docx"; _docx(p)
    ex = extract_document(p)
    assert ex.content_type == "document" and ex.title == "laudo final" and ex.channel == "documento docx"
    assert "## Laudo de Avaliação" in ex.text and "### Metodologia" in ex.text
    assert "| Item | Valor |" in ex.text and "| Terreno | 1000 |" in ex.text


def test_xlsx_sheets_to_markdown(tmp_path):
    p = tmp_path / "custos.xlsx"; _xlsx(p)
    ex = extract_document(p)
    assert "## Aba: Custos" in ex.text and "| Servidor | 120.5 |" in ex.text and "Vazia" not in ex.text


def test_pdf_without_text_and_unsupported(tmp_path):
    from pypdf import PdfWriter
    w = PdfWriter(); w.add_blank_page(width=200, height=200); w.write(str(tmp_path / "scan.pdf"))
    with pytest.raises(ExtractionError) as e:
        extract_document(tmp_path / "scan.pdf")
    assert e.value.code == "ERR-004" and not e.value.retryable
    (tmp_path / "x.exe").write_bytes(b"x" * 300)
    with pytest.raises(ExtractionError) as e2:
        extract_document(tmp_path / "x.exe")
    assert e2.value.code == "ERR-002"


def test_normalize_file_url_hashes_content(tmp_path):
    a = tmp_path / "a.txt"; a.write_text("conteudo" * 50)
    b = tmp_path / "b.txt"; b.write_text("conteudo" * 50)
    na, nb = normalize(f"file://{a}"), normalize(f"file://{b}")
    assert na.platform == "document" and na.external_id == nb.external_id  # mesmo conteúdo, nomes diferentes → duplicata
    assert normalize(f"file://{tmp_path}/nao-existe.pdf").external_id is None


def test_vault_document_item_has_full_text(tmp_path, monkeypatch):
    import json
    monkeypatch.setattr(vault.settings, "vault_dir", tmp_path)
    e = {"title": "Laudo", "summary": "S.", "key_points": ["a"], "tags": ["laudo"], "confidence": "alta",
         "applicability": {"saas_pessoal": "nenhuma", "projeto_cliente": "alta", "estudo_geral": "media"}, "how_to_apply": ""}
    item = {"id": "1", "platform": "document", "external_id": "abc", "canonical_url": "file:///app/tmp/documents/xyz.docx",
            "channel": "documento docx", "captured_at": "2026-08-27T10:00:00+00:00", "status": "enriched", "triage_decision": None,
            "content_type": "document", "vault_path": None, "title": "Laudo", "raw_content": "# Laudo\n\ntexto", "enrichment": json.dumps(e)}
    md = vault.write_item(item).read_text()
    assert "📎 xyz.docx (enviado pelo Telegram)" in md and "## Texto integral" in md and "file://" not in md


def test_capture_line_noun():
    import random
    for sat in voices.CAPTURE:
        for i in range(4):
            line = voices.capture_line(1, sat=sat, noun="arquivo", rng=random.Random(i))
            assert "{" not in line and "link" not in line
